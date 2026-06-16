import os
import sys
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Configuration des chemins
BASE_DIR = "/home/tresor/B2B"
SCRATCH_DIR = os.path.join(BASE_DIR, "scratch")
IMAGES_DIR = os.path.join(SCRATCH_DIR, "images")
OUTPUT_DOCX = os.path.join(BASE_DIR, "rapport_projet_b2b.docx")

os.makedirs(IMAGES_DIR, exist_ok=True)

FONT_REGULAR = "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"

# ==========================================
# 1. OUTILS DE DESSIN PIL (DIAGRAMMES)
# ==========================================

def get_fonts(size):
    try:
        reg = ImageFont.truetype(FONT_REGULAR, size)
        bld = ImageFont.truetype(FONT_BOLD, size)
        return reg, bld
    except Exception as e:
        print(f"Erreur chargement polices: {e}. Utilisation police par défaut.")
        return ImageFont.load_default(), ImageFont.load_default()

def draw_actor(draw, x, y, name, font, color):
    draw.ellipse([x-15, y-50, x+15, y-20], outline=color, width=2)
    draw.line([x, y-20, x, y+20], fill=color, width=2)
    draw.line([x-25, y-5, x+25, y-5], fill=color, width=2)
    draw.line([x, y+20, x-15, y+55], fill=color, width=2)
    draw.line([x, y+20, x+15, y+55], fill=color, width=2)
    try:
        w = int(draw.textlength(name, font=font))
    except:
        w = len(name)*6
    draw.text((x - w//2, y+60), name, fill=color, font=font)

def draw_ellipse_uc(draw, x, y, w, h, text, font, border_color, fill_color, text_color):
    draw.ellipse([x - w//2, y - h//2, x + w//2, y + h//2], fill=fill_color, outline=border_color, width=2)
    lines = text.split('\n')
    total_h = len(lines) * 16
    start_y = y - total_h // 2
    for i, line in enumerate(lines):
        try:
            lw = int(draw.textlength(line, font=font))
        except:
            lw = len(line)*6
        draw.text((x - lw//2, start_y + i*16), line, fill=text_color, font=font)

def draw_class_box(draw, x, y, w, title, attributes, methods, font_title, font_body, border_color, title_bg, text_color):
    header_h = 32
    draw.rectangle([x, y, x+w, y+header_h], fill=title_bg, outline=border_color, width=2)
    try:
        tw = int(draw.textlength(title, font=font_title))
    except:
        tw = len(title)*7
    draw.text((x + (w-tw)//2, y + 6), title, fill=text_color, font=font_title)
    
    body_h = len(attributes)*20 + len(methods)*20 + 20
    draw.rectangle([x, y+header_h, x+w, y+header_h+body_h], fill="#FFFFFF", outline=border_color, width=2)
    
    curr_y = y + header_h + 8
    for attr in attributes:
        draw.text((x + 10, curr_y), attr, fill="#333333", font=font_body)
        curr_y += 20
        
    if methods:
        draw.line([x, curr_y - 4, x+w, curr_y - 4], fill=border_color, width=1)
        for m in methods:
            draw.text((x + 10, curr_y), m, fill="#333333", font=font_body)
            curr_y += 20
            
    return y + header_h + body_h

def draw_lifeline(draw, x, y_start, y_end, name, font, color):
    w, h = 150, 44
    draw.rectangle([x - w//2, y_start - h//2, x + w//2, y_start + h//2], fill="#1B365D", outline="#11223F", width=2)
    try:
        tw = int(draw.textlength(name, font=font))
    except:
        tw = len(name)*7
    draw.text((x - tw//2, y_start - 8), name, fill="#FFFFFF", font=font)
    for curr_y in range(y_start + h//2 + 5, y_end, 15):
        draw.line([x, curr_y, x, min(curr_y + 8, y_end)], fill=color, width=1)

def draw_msg(draw, x_from, x_to, y, label, font, is_dashed=False, is_return=False, color="#1B365D"):
    if is_dashed:
        for curr_x in range(min(x_from, x_to), max(x_from, x_to), 10):
            draw.line([curr_x, y, min(curr_x + 6, max(x_from, x_to)), y], fill=color, width=1)
    else:
        draw.line([x_from, y, x_to, y], fill=color, width=2)
        
    arrow_size = 8
    direction = 1 if x_to > x_from else -1
    if is_return:
        draw.line([x_to, y, x_to - direction*arrow_size, y - arrow_size//2], fill=color, width=2)
        draw.line([x_to, y, x_to - direction*arrow_size, y + arrow_size//2], fill=color, width=2)
    else:
        draw.polygon([
            (x_to, y),
            (x_to - direction*arrow_size, y - arrow_size//2),
            (x_to - direction*arrow_size, y + arrow_size//2)
        ], fill=color)
        
    try:
        tw = int(draw.textlength(label, font=font))
    except:
        tw = len(label)*6
    text_x = min(x_from, x_to) + abs(x_to - x_from)//2 - tw//2
    draw.text((text_x, y - 18), label, fill="#333333", font=font)

# ------------------------------------------
# Génération de chaque diagramme en PNG
# ------------------------------------------

def generate_diagrams():
    reg12, bld12 = get_fonts(12)
    reg14, bld14 = get_fonts(14)
    reg10, bld10 = get_fonts(10)

    # 1. Use Case Diagram
    img = Image.new("RGB", (1000, 750), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1000, 750], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE CAS D'UTILISATION - MARKETPLACE B2B", fill="#1B365D", font=bld14)
    
    draw.rectangle([250, 80, 750, 700], fill="#FFFFFF", outline="#1B365D", width=2)
    draw.text((260, 90), "Système B2B (Spring Boot + React)", fill="#555555", font=reg12)
    
    draw_actor(draw, 120, 350, "Visiteur / Acheteur", bld12, "#1B365D")
    draw_actor(draw, 880, 380, "Vendeur (Authentifié)", bld12, "#1B365D")
    
    ucs = [
        (500, 140, 220, 50, "S'enregistrer / Se connecter"),
        (500, 220, 240, 50, "Rechercher / Filtrer les produits"),
        (500, 300, 240, 50, "Consulter les détails d'un produit"),
        (500, 380, 240, 50, "Gérer les favoris (Ajout/Suivi)"),
        (500, 460, 260, 50, "Contacter le vendeur (WhatsApp Direct)"),
        (500, 540, 260, 50, "Partager l'annonce avec un ami"),
        (500, 620, 240, 50, "Créer / Éditer / Supprimer Produit")
    ]
    
    for x, y, w, h, txt in ucs:
        draw_ellipse_uc(draw, x, y, w, h, txt, reg12, "#1B365D", "#E3F2FD", "#1B365D")
        
    for i in [0, 1, 2, 3, 4, 5]:
        draw.line([120, 350, ucs[i][0]-ucs[i][2]//2, ucs[i][1]], fill="#2C3E50", width=1)
    for i in [0, 2, 3, 5, 6]:
        draw.line([880, 380, ucs[i][0]+ucs[i][2]//2, ucs[i][1]], fill="#2C3E50", width=1)
        
    img.save(os.path.join(IMAGES_DIR, "use_case_diagram.png"), "PNG")
    print("Généré: use_case_diagram.png")

    # 2. Business Class Diagram
    img = Image.new("RGB", (1000, 550), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1000, 550], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE CLASSES METIER (DOMAIN MODEL)", fill="#1B365D", font=bld14)
    
    draw_class_box(draw, 60, 130, 240, "User", 
                   ["- id: Long", "- name: String", "- email: String", "- password: String", "- whatsapp: String"],
                   ["+ getProduits(): List<Produit>", "+ getFavoris(): Set<Produit>"],
                   bld12, reg12, "#1B365D", "#1B365D", "#FFFFFF")
                   
    draw_class_box(draw, 420, 100, 270, "Produit", 
                   ["- id: Long", "- name: String", "- description: String", "- price: Long", "- views: int", "- statut: StatutProduit", "- brand: String", "- whatsapp: String", "- nouveau: boolean", "- misEnAvant: boolean", "- image: byte[]", "- datePublication: LocalDate"],
                   ["+ getVendeur(): User", "+ getCategorie(): CategorieProduit"],
                   bld12, reg12, "#1B365D", "#1B365D", "#FFFFFF")
                   
    draw_class_box(draw, 780, 170, 180, "CategorieProduit", 
                   ["- id: Long", "- name: String"],
                   ["+ getProduits(): List<Produit>"],
                   bld12, reg12, "#1B365D", "#1B365D", "#FFFFFF")
                   
    # Relations
    draw.line([300, 200, 420, 200], fill="#1B365D", width=2)
    draw.text((310, 180), "1", fill="#333333", font=reg12)
    draw.text((380, 180), "0..*", fill="#333333", font=reg12)
    draw.text((320, 205), "publie / vend", fill="#555555", font=reg10)
    
    draw.line([690, 240, 780, 240], fill="#1B365D", width=2)
    draw.text((700, 220), "0..*", fill="#333333", font=reg12)
    draw.text((760, 220), "1", fill="#333333", font=reg12)
    draw.text((705, 245), "appartient à", fill="#555555", font=reg10)
    
    draw.line([180, 310, 180, 450], fill="#2980B9", width=2)
    draw.line([180, 450, 550, 450], fill="#2980B9", width=2)
    draw.line([550, 450, 550, 410], fill="#2980B9", width=2)
    draw.text((190, 425), "0..*", fill="#333333", font=reg12)
    draw.text((510, 425), "0..*", fill="#333333", font=reg12)
    draw.text((320, 455), "a en favoris (M2M)", fill="#555555", font=reg10)

    img.save(os.path.join(IMAGES_DIR, "business_class_diagram.png"), "PNG")
    print("Généré: business_class_diagram.png")

    # 3. Technical Class Diagram
    img = Image.new("RGB", (1100, 700), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1100, 700], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE CLASSES TECHNIQUE (ARCHITECTURE EN COUCHES)", fill="#1B365D", font=bld14)
    
    draw.text((40, 110), "COUCHES CONTROLEURS (REST API)", fill="#2C3E50", font=bld12)
    draw.text((40, 270), "COUCHES SERVICES (LOGIQUE METIER)", fill="#2C3E50", font=bld12)
    draw.text((40, 430), "COUCHES REPOSITORIES (ACCES DONNEES)", fill="#2C3E50", font=bld12)
    draw.text((40, 570), "COUCHES ENTITES (ENTITIES & DTOS)", fill="#2C3E50", font=bld12)
    
    draw_class_box(draw, 350, 90, 250, "ProduitController",
                   ["+ findAll(): ResponseEntity", "+ findById(): ResponseEntity", "+ create(): ResponseEntity", "+ update(): ResponseEntity"],
                   [], bld10, reg10, "#1B365D", "#1B365D", "#FFFFFF")
    draw_class_box(draw, 700, 90, 250, "UserController",
                   ["+ register(): ResponseEntity", "+ login(): ResponseEntity", "+ toggleFavorite(): ResponseEntity"],
                   [], bld10, reg10, "#1B365D", "#1B365D", "#FFFFFF")
                   
    draw_class_box(draw, 350, 250, 250, "ProduitService",
                   ["- repository: ProduitRepository", "- userService: UserService"],
                   ["+ findAll(): List<Produit>", "+ create(): Produit", "+ update(): Produit"],
                   bld10, reg10, "#16A085", "#16A085", "#FFFFFF")
    draw_class_box(draw, 700, 250, 250, "UserService",
                   ["- repository: UserRepository", "- encoder: PasswordEncoder"],
                   ["+ register(): User", "+ login(): User", "+ toggleFavorite(): boolean"],
                   bld10, reg10, "#16A085", "#16A085", "#FFFFFF")
                   
    draw_class_box(draw, 350, 410, 250, "ProduitRepository",
                   ["<<interface>>", "extends JpaRepository<Produit, Long>"],
                   ["+ findByVendeurId()", "+ findFeatured()"],
                   bld10, reg10, "#D35400", "#D35400", "#FFFFFF")
    draw_class_box(draw, 700, 410, 250, "UserRepository",
                   ["<<interface>>", "extends JpaRepository<User, Long>"],
                   ["+ findByEmail()"],
                   bld10, reg10, "#D35400", "#D35400", "#FFFFFF")
                   
    draw_class_box(draw, 380, 560, 180, "Produit",
                   ["- id: Long", "- name: String", "- price: Long"], [],
                   bld10, reg10, "#7F8C8D", "#7F8C8D", "#FFFFFF")
    draw_class_box(draw, 730, 560, 180, "User",
                   ["- id: Long", "- name: String", "- email: String"], [],
                   bld10, reg10, "#7F8C8D", "#7F8C8D", "#FFFFFF")
                   
    draw.line([475, 202, 475, 250], fill="#333333", width=2)
    draw.line([825, 202, 825, 250], fill="#333333", width=2)
    draw.line([475, 362, 475, 410], fill="#333333", width=2)
    draw.line([825, 362, 825, 410], fill="#333333", width=2)
    draw.line([475, 482, 475, 560], fill="#333333", width=1)
    draw.line([825, 482, 825, 560], fill="#333333", width=1)
    
    draw.line([600, 300, 700, 300], fill="#E74C3C", width=2)
    draw.text((615, 280), "calls", fill="#E74C3C", font=reg10)

    img.save(os.path.join(IMAGES_DIR, "technical_class_diagram.png"), "PNG")
    print("Généré: technical_class_diagram.png")

    # 4. Sequence Diagram: Add Product
    img = Image.new("RGB", (1000, 750), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1000, 750], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE SEQUENCE : AJOUT DE PRODUIT", fill="#1B365D", font=bld14)
    
    draw_lifeline(draw, 150, 100, 700, "React UI", bld12, "#1B365D")
    draw_lifeline(draw, 360, 100, 700, "ProduitController", bld12, "#1B365D")
    draw_lifeline(draw, 570, 100, 700, "ProduitService", bld12, "#1B365D")
    draw_lifeline(draw, 780, 100, 700, "ProduitRepository", bld12, "#1B365D")
    draw_lifeline(draw, 920, 100, 700, "Neon DB", bld12, "#1B365D")
    
    draw_msg(draw, 150, 360, 170, "POST /api/produits (Form-Data + JWT)", reg10)
    draw.rectangle([340, 195, 380, 225], fill="#BDC3C7", outline="#7F8C8D")
    draw.text((345, 205), "JWT OK", fill="#1B365D", font=bld10)
    
    draw_msg(draw, 360, 570, 270, "create(vendeurId, name, price, ...)", reg10)
    
    draw_msg(draw, 570, 780, 340, "findById() / findByNom()", reg10)
    draw_msg(draw, 780, 920, 370, "SELECT ...", reg10)
    draw_msg(draw, 920, 780, 420, "User/Categorie Entities", reg10, is_dashed=True, is_return=True)
    draw_msg(draw, 780, 570, 450, "Return Entities", reg10, is_dashed=True, is_return=True)
    
    draw_msg(draw, 570, 780, 510, "save(produit)", reg10)
    draw_msg(draw, 780, 920, 540, "INSERT INTO produits ...", reg10)
    draw_msg(draw, 920, 780, 590, "INSERT SUCCESS", reg10, is_dashed=True, is_return=True)
    draw_msg(draw, 780, 570, 620, "Saved Produit entity", reg10, is_dashed=True, is_return=True)
    
    draw_msg(draw, 570, 360, 650, "Produit entity", reg10, is_dashed=True, is_return=True)
    draw_msg(draw, 360, 150, 680, "201 Created (ProduitResponse DTO)", reg10, is_dashed=True, is_return=True)
    
    img.save(os.path.join(IMAGES_DIR, "sequence_add_product.png"), "PNG")
    print("Généré: sequence_add_product.png")

    # 5. Sequence Diagram: Contact Seller
    img = Image.new("RGB", (1000, 650), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1000, 650], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE SEQUENCE : SE METTRE EN CONTACT AVEC UN VENDEUR", fill="#1B365D", font=bld14)
    
    draw_lifeline(draw, 150, 100, 600, "Acheteur (User)", bld12, "#1B365D")
    draw_lifeline(draw, 400, 100, 600, "React Frontend", bld12, "#1B365D")
    draw_lifeline(draw, 680, 100, 600, "Spring Boot API", bld12, "#1B365D")
    draw_lifeline(draw, 900, 100, 600, "WhatsApp Web/App", bld12, "#1B365D")
    
    draw_msg(draw, 150, 400, 170, "Consulter Détails Produit X", reg10)
    draw_msg(draw, 400, 680, 230, "PATCH /api/produits/{id}/vues (incrémenter)", reg10)
    draw_msg(draw, 680, 400, 290, "200 OK (Produit details avec whatsapp)", reg10, is_dashed=True, is_return=True)
    
    draw_msg(draw, 150, 400, 360, "Clic bouton 'Contacter le Vendeur'", reg10)
    draw.rectangle([350, 395, 450, 425], fill="#E8F8F5", outline="#2ECC71")
    draw.text((360, 405), "Générer URL wa.me", fill="#16A085", font=bld10)
    
    draw_msg(draw, 400, 150, 470, "Redirection du navigateur", reg10, is_dashed=True, is_return=True)
    draw_msg(draw, 150, 900, 520, "Ouvrir chat avec texte prérempli", reg10)
    draw_msg(draw, 900, 150, 570, "Envoi message (Bonjour, je suis intéressé par...)", reg10, is_dashed=True, is_return=True)
    
    img.save(os.path.join(IMAGES_DIR, "sequence_contact_seller.png"), "PNG")
    print("Généré: sequence_contact_seller.png")

    # 6. Sequence Diagram: Share Ad
    img = Image.new("RGB", (1000, 650), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 0, 1000, 650], fill="#F4F6F9", outline="#D3D3D3", width=2)
    draw.text((50, 25), "DIAGRAMME DE SEQUENCE : PARTAGE DE L'ANNONCE AVEC UN AMI", fill="#1B365D", font=bld14)
    
    draw_lifeline(draw, 150, 100, 600, "Utilisateur (Partageur)", bld12, "#1B365D")
    draw_lifeline(draw, 400, 100, 600, "React App UI", bld12, "#1B365D")
    draw_lifeline(draw, 680, 100, 600, "Navigateur (Web Share)", bld12, "#1B365D")
    draw_lifeline(draw, 900, 100, 600, "Ami (Destinataire)", bld12, "#1B365D")
    
    draw_msg(draw, 150, 400, 170, "Clic sur bouton 'Partager l'annonce'", reg10)
    
    draw.rectangle([350, 205, 450, 235], fill="#EBF5FB", outline="#3498DB")
    draw.text((355, 215), "Check Web Share API", fill="#2980B9", font=bld10)
    
    draw_msg(draw, 400, 680, 290, "navigator.share({title, text, url})", reg10)
    draw_msg(draw, 680, 150, 350, "Afficher boite de dialogue native", reg10, is_dashed=True, is_return=True)
    
    draw_msg(draw, 150, 680, 420, "Sélectionner Ami via WhatsApp/Mail/SMS", reg10)
    draw_msg(draw, 680, 900, 480, "Envoi du lien URL de l'annonce", reg10)
    draw_msg(draw, 900, 400, 540, "Clic sur l'URL & consulte l'annonce", reg10)
    
    img.save(os.path.join(IMAGES_DIR, "sequence_share_ad.png"), "PNG")
    print("Généré: sequence_share_ad.png")

# ==========================================
# 2. COMPILATION DU DOCUMENT DOCX
# ==========================================

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_callout(doc, text, title="NOTE", color="1B365D"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.0)
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F0F4F8")
    
    # Set left border only
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    
    run_title = p.add_run(f"[{title}] ")
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    run_title.font.name = "Arial"
    run_title.font.size = Pt(10)
    
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_para(p, size=11, bold=False, italic=False, color=None, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = color

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.keep_with_next = True
    
    if level == 1:
        color = RGBColor(0x1B, 0x36, 0x5D)
        size = 18
        space_before = 18
        space_after = 8
        pPr = h._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="4" w:color="1B365D"/></w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        color = RGBColor(0x2C, 0x3E, 0x50)
        size = 14
        space_before = 14
        space_after = 6
    else:
        color = RGBColor(0x55, 0x55, 0x55)
        size = 12
        space_before = 10
        space_after = 4
        
    for r in h.runs:
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = True
        r.font.color.rgb = color
        
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    return h

def generate_report():
    doc = Document()
    
    # Page margins
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    footer = sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    frun = fp.add_run("Rapport de Projet — Cours de Conception SI  |  Marketplace B2B  |  Page ")
    frun.font.name = "Arial"
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    add_page_number(fp.add_run())
    
    header = sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hrun = hp.add_run("Cours de Conception des Systèmes d'Information — Architecture, Modélisation UML et Déploiement")
    hrun.font.name = "Arial"
    hrun.font.size = Pt(9)
    hrun.font.italic = True
    hrun.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    hpPr = hp._p.get_or_add_pPr()
    hpBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="1" w:color="BDC3C7"/></w:pBdr>')
    hpPr.append(hpBdr)

    # PAGE DE GARDE
    p_space = doc.add_paragraph()
    format_para(p_space, space_before=40, space_after=0)

    # Établissement
    p_school = doc.add_paragraph()
    p_school.add_run("COURS DE CONCEPTION DES SYSTÈMES D'INFORMATION")
    format_para(p_school, size=13, bold=True, color=RGBColor(0x7F, 0x8C, 0x8D), space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_dept = doc.add_paragraph()
    p_dept.add_run("Filière Informatique — Année Académique 2025-2026")
    format_para(p_dept, size=11, italic=True, color=RGBColor(0x7F, 0x8C, 0x8D), space_after=20, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_divider0 = doc.add_paragraph()
    p_divider0.add_run("─" * 50)
    format_para(p_divider0, size=11, color=RGBColor(0xBD, 0xC3, 0xC7), space_after=25, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_title_pre = doc.add_paragraph()
    p_title_pre.add_run("PLATEFORME MARKETPLACE B2B FULLSTACK")
    format_para(p_title_pre, size=14, bold=True, color=RGBColor(0x7F, 0x8C, 0x8D), space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_title = doc.add_paragraph()
    p_title.add_run("RAPPORT DE PROJET :\nCONCEPTION, RÉALISATION ET DÉPLOIEMENT")
    format_para(p_title, size=22, bold=True, color=RGBColor(0x1B, 0x36, 0x5D), space_after=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_divider = doc.add_paragraph()
    p_divider.add_run("—" * 45)
    format_para(p_divider, size=12, bold=True, color=RGBColor(0x1B, 0x36, 0x5D), space_after=25, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.add_run(
        "Ce rapport présente la conception et la réalisation d'une plateforme web d'annonces interentreprises (B2B) fullstack, "
        "réalisée dans le cadre du cours de Conception des Systèmes d'Information. "
        "Il couvre l'analyse fonctionnelle (cas d'utilisation), la modélisation UML (diagrammes de classes et de séquences), "
        "le guide d'utilisation, ainsi que l'architecture de déploiement adoptée."
    )
    format_para(p_subtitle, size=11, italic=True, color=RGBColor(0x34, 0x49, 0x5E), space_after=35, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_details = doc.add_paragraph()
    p_details.add_run("Réalisé par : [Prénom NOM 1]  —  [Prénom NOM 2]\n")
    p_details.add_run("Encadrant(e) : [Titre Prénom NOM du Professeur]\n")
    p_details.add_run("Technologies : Spring Boot (Java 17) | React Vite | PostgreSQL Neon | Docker\n")
    p_details.add_run("Déploiement : Backend AWS EC2  •  Frontend Vercel  •  Base de données Neon\n")
    p_details.add_run("Date de soumission : Juin 2026")
    format_para(p_details, size=11, bold=True, color=RGBColor(0x2C, 0x3E, 0x50), space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()

    # SOMMAIRE
    add_heading_styled(doc, "TABLE DES MATIÈRES", level=1)
    p_toc_desc = doc.add_paragraph()
    p_toc_desc.add_run("Ce sommaire présente la structure globale du document technique. "
                       "Les sections et numéros de page correspondent à l'organisation séquentielle de ce rapport de projet.")
    format_para(p_toc_desc, italic=True, space_after=15)
    
    toc_items = [
        ("INTRODUCTION ET CONTEXTE DU PROJET", "3"),
        ("1. Contexte du Projet et Cours de Conception SI", "3"),
        ("2. Problématique et Objectifs", "3"),
        ("3. Présentation Générale de l'Application Réalisée", "4"),
        ("CAHIER DES CHARGES ET SPÉCIFICATIONS FONCTIONNELLES", "5"),
        ("1. Acteurs du Système et Droits d'Accès", "5"),
        ("2. Spécifications Fonctionnelles par Cas d'Utilisation", "5"),
        ("3. Diagramme de Cas d'Utilisation Global du Système", "6"),
        ("ARCHITECTURE SYSTÈME ET CHOIX TECHNOLOGIQUES", "7"),
        ("1. Architecture Globale et Modèle Multi-tiers", "7"),
        ("2. Backend : Spring Boot (Java 17)", "7"),
        ("3. Frontend : React avec Vite", "8"),
        ("4. Base de Données : PostgreSQL Neon", "8"),
        ("5. Containerisation avec Docker", "9"),
        ("CONCEPTION DU SYSTÈME & MODÉLISATION UML", "10"),
        ("1. Diagramme de Classes Métier (Domain Model)", "10"),
        ("2. Diagramme de Classes Technique (Architecture en Couches)", "11"),
        ("3. Dictionnaire de Données et Structure des Tables SQL", "12"),
        ("MODÈLE DE SÉCURITÉ ET AUTHENTIFICATION JWT", "13"),
        ("1. Principes de la Sécurité Stateless avec JWT", "13"),
        ("2. Filtre de Sécurité dans Spring Boot", "13"),
        ("3. Gestion de la Politique CORS", "14"),
        ("CONCEPTION DÉTAILLÉE DES FLUX ET SÉQUENCES", "15"),
        ("1. Séquence : Publication d'un Produit par le Vendeur", "15"),
        ("2. Séquence : Contact Vendeur via WhatsApp", "16"),
        ("3. Séquence : Partage d'une Annonce avec un Ami", "17"),
        ("4. Implémentation du Partage (Web Share API)", "17"),
        ("DEVOPS ET WORKFLOW DE DÉPLOIEMENT AVEC DOCKER", "19"),
        ("1. Containerisation du Backend avec Dockerfile", "19"),
        ("2. Orchestration Locale avec Docker Compose", "19"),
        ("ARCHITECTURE DE DÉPLOIEMENT CLOUD", "21"),
        ("1. Vue d'Ensemble de l'Architecture Déployée", "21"),
        ("2. Hébergement Backend sur AWS EC2", "21"),
        ("3. Hébergement Frontend sur Vercel", "22"),
        ("4. Communication entre les Composants", "22"),
        ("GUIDE D'UTILISATION DE LA PLATEFORME B2B", "23"),
        ("1. Inscription, Connexion et Gestion du Profil", "23"),
        ("2. Recherche, Filtres par Catégorie et Favoris", "23"),
        ("3. Espace Vendeur : Publier, Modifier, Supprimer une Annonce", "24"),
        ("4. Contact Vendeur et Partage d'Annonces", "24"),
        ("VALIDATION ET TESTS DE L'APPLICATION", "26"),
        ("1. Tests Fonctionnels des Scénarios Clés", "26"),
        ("2. Résultats et Points d'Amélioration", "26"),
        ("CONCLUSION ET PERSPECTIVES", "27")
    ]
    
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        is_main = title.isupper()
        indent = Inches(0) if is_main else Inches(0.3)
        p.paragraph_format.left_indent = indent
        
        run_title = p.add_run(title)
        run_title.font.name = "Arial"
        run_title.font.size = Pt(10 if is_main else 9.5)
        run_title.bold = is_main
        if is_main:
            run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            
        dots_count = 80 - len(title) - len(page)
        if not is_main:
            dots_count -= 6
        dots_count = max(5, dots_count)
        
        run_dots = p.add_run(" ." * dots_count + " ")
        run_dots.font.name = "Arial"
        run_dots.font.size = Pt(9)
        run_dots.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
        
        run_page = p.add_run(page)
        run_page.font.name = "Arial"
        run_page.font.size = Pt(10 if is_main else 9.5)
        run_page.bold = is_main
        if is_main:
            run_page.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            
    doc.add_page_break()

    # CHAPITRE 1
    add_heading_styled(doc, "INTRODUCTION ET CONTEXTE DU PROJET", level=1)
    
    add_heading_styled(doc, "1. Contexte du Projet et Cours de Conception SI", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Ce projet a été réalisé dans le cadre du cours de Conception des Systèmes d'Information. "
        "L'objectif est de concevoir et de développer une application web fullstack de type marketplace B2B "
        "(Business-to-Business), c'est-à-dire une plateforme d'annonces en ligne dédiée aux échanges entre professionnels. "
        "Le système permet à des vendeurs de publier des annonces de produits avec photos, et aux acheteurs de les consulter, "
        "de les mettre en favoris, de contacter directement le vendeur via WhatsApp, et de partager une annonce avec un tiers. "
        "L'application est construite avec une architecture multicouche (frontend, backend, base de données) et est déployée "
        "sur une infrastructure cloud distribuée."
    )
    format_para(p, space_after=12)
    
    add_heading_styled(doc, "2. Problématique et Objectifs", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La problématique de ce projet est la suivante : comment concevoir et mettre en œuvre une plateforme web "
        "B2B permettant à des professionnels de publier, consulter et partager des annonces, avec une architecture "
        "sécurisée, modulaire et déployable sur le cloud ? "
        "Les objectifs fixés pour le projet sont :"
        "\n- Sécurité des accès : Seuls les utilisateurs authentifiés peuvent publier ou modifier des annonces. "
        "La consultation du catalogue reste accessible à tous les visiteurs."
        "\n- Communication directe : Intégrer un mécanisme de contact via WhatsApp sans passer par un système de messagerie interne."
        "\n- Partage d'annonces : Permettre à un utilisateur de partager un lien d'annonce via l'API Web Share du navigateur "
        "ou par copie dans le presse-papiers."
        "\n- Déploiement cloud : L'application est hébergée sur une infrastructure cloud réelle "
        "(backend sur AWS EC2, frontend sur Vercel, base de données sur Neon DB)."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Présentation Générale de l'Application Réalisée", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'application est un portail web fullstack organisé en deux espaces. "
        "L'espace public est accessible à tout visiteur sans authentification : il permet de consulter le catalogue de produits, "
        "d'effectuer des recherches par mots-clés, de filtrer par catégories, et de lire les fiches détaillées des annonces. "
        "L'espace privé, accessible aux utilisateurs connectés, donne accès au tableau de bord vendeur "
        "(publication, modification et suppression d'annonces), à la liste personnelle de favoris, "
        "ainsi qu'aux fonctionnalités de contact et de partage d'annonces. "
        "L'interface est construite en React (Vite). Le backend Spring Boot (Java 17) expose une API REST consommée par le frontend. "
        "La base de données PostgreSQL hébergée sur Neon stocke les utilisateurs, les produits et les associations de favoris."
    )
    format_para(p, space_after=12)
    
    add_callout(doc, 
                "La plateforme développée couvre l'ensemble du cycle d'une application web moderne : "
                "modélisation du système, développement backend et frontend, sécurisation par JWT, "
                "et déploiement sur une infrastructure cloud distribuée (AWS EC2, Vercel, Neon DB).",
                "PÉRIMÈTRE DU PROJET", "1B365D")

    doc.add_page_break()

    # CHAPITRE 2
    add_heading_styled(doc, "CAHIER DES CHARGES ET SPÉCIFICATIONS FONCTIONNELLES", level=1)
    
    add_heading_styled(doc, "1. Acteurs du Système et Droits d'Accès", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'analyse fonctionnelle de la plateforme identifie trois catégories d'acteurs principaux qui interagissent avec le système. "
        "Chacun possède des droits d'accès et des fonctionnalités spécifiques qui garantissent la sécurité et la cohérence de l'utilisation :"
    )
    format_para(p, space_after=8)
    
    actors = [
        ("Visiteur / Acheteur non-authentifié", 
         "Cet acteur représente le point d'entrée sur l'application. Sans avoir besoin de s'authentifier, "
         "il a le droit de lire le catalogue global. Il peut effectuer des recherches par mots-clés dans le moteur de recherche interne, "
         "filtrer le catalogue par catégories prédéfinies, et ouvrir la page de détails d'une annonce. C'est sur cette page de détails qu'il peut "
         "cliquer sur le bouton vert pour contacter directement le vendeur par WhatsApp (avec pré-remplissage du texte du message) "
         "ou utiliser l'icône de partage pour envoyer le lien de la publication à un ami ou un collègue."),
        ("Acheteur authentifié", 
         "Cet acteur a créé un compte utilisateur sur la plateforme et s'est authentifié à l'aide de ses identifiants. "
         "En plus des droits conférés au visiteur anonyme, l'acheteur authentifié peut mettre en favoris les annonces qui l'intéressent. "
         "Un bouton en forme de cœur sur chaque carte de produit lui permet d'ajouter ou de retirer un produit de ses favoris. "
         "Il dispose d'une page dédiée (Favorites.jsx) qui filtre et affiche en temps réel la liste de tous ses produits favoris, "
         "lui permettant de suivre de près l'évolution des offres."),
        ("Vendeur professionnel", 
         "Le rôle de vendeur est accessible à tout utilisateur authentifié qui souhaite proposer des produits ou matériels sur la plateforme. "
         "Il accède à un espace de gestion appelé Dashboard (Dashboard.jsx). De là, il peut soumettre de nouveaux produits via un formulaire multipart "
         "(titre, marque, prix, description textuelle, photo du produit et numéro WhatsApp spécifique pour l'annonce). Le vendeur est également "
         "le seul habilité à éditer les informations de ses propres produits, à modifier la photo, à supprimer l'annonce si le matériel n'est plus "
         "disponible, ou à modifier son statut pour le marquer comme 'Vendu'.")
    ]
    
    for act_name, act_desc in actors:
        p_act = doc.add_paragraph()
        r_name = p_act.add_run(f"• {act_name} : ")
        r_name.bold = True
        r_name.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        p_act.add_run(act_desc)
        format_para(p_act, space_after=6)

    add_heading_styled(doc, "2. Spécifications Fonctionnelles Détaillées par Cas d'Utilisation", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Les cas d'utilisation modélisent le comportement attendu du système en réponse aux actions des acteurs. "
        "Voici les spécifications comportementales pour chaque cas clé du cahier des charges :"
        "\n- Enregistrement & Authentification : L'utilisateur doit pouvoir soumettre son nom, son email, son mot de passe et son numéro de téléphone WhatsApp. "
        "Le mot de passe doit être chiffré à l'aide d'un algorithme robuste côté serveur avant l'écriture en base de données. En retour, un jeton d'accès sécurisé "
        "JWT est généré et stocké dans le stockage local du navigateur client (localStorage)."
        "\n- Publication d'une annonce : Le vendeur remplit les détails de l'annonce et téléverse une image au format standard (JPEG/PNG). "
        "Le système valide que l'utilisateur est authentifié grâce au JWT avant de stocker l'image et les métadonnées dans la base de données. "
        "Le produit nouvellement créé prend par défaut le statut 'Disponible'."
        "\n- Prise de contact : L'acheteur clique sur 'Contacter via WhatsApp'. L'application détecte si le numéro WhatsApp du produit est fourni. "
        "Elle génère une adresse URL universelle WhatsApp intégrant le texte du message et redirige l'acheteur. Ce processus assure "
        "une communication fluide et directe sans intermédiaire de paiement."
        "\n- Partage d'annonces : L'utilisateur clique sur le bouton de partage. Si l'API navigator.share est supportée, l'appareil affiche son volet "
        "de partage système pour transmettre le lien de l'annonce. Sinon, le lien est copié dans le presse-papiers, permettant un partage manuel instantané."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Diagramme de Cas d'Utilisation Global du Système", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le diagramme ci-dessous schématise l'ensemble des interactions entre les différents types d'utilisateurs "
        "(Acheteur, Visiteur et Vendeur) et le système central de la marketplace B2B. Ce diagramme structure graphiquement "
        "les frontières de l'application et les relations d'inclusion/extension entre cas d'utilisation :"
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "use_case_diagram.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 2.1 : Diagramme de Cas d'Utilisation Global du Système B2B")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)
    
    doc.add_page_break()

    # CHAPITRE 3
    add_heading_styled(doc, "ARCHITECTURE SYSTÈME ET CHOIX TECHNOLOGIQUES", level=1)
    
    add_heading_styled(doc, "1. Architecture Globale et Modèle Multi-tiers", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'application repose sur un modèle d'architecture découplé de type multi-tiers (3-Tier Architecture), "
        "assurant une séparation stricte des responsabilités et une scalabilité optimale. Ce modèle comprend :"
        "\n- Le Tier de Présentation (Frontend) : Construit sous forme de Single Page Application (SPA) en React. "
        "Ce client léger interagit de manière asynchrone avec le backend via des requêtes HTTP REST. Le fait d'utiliser une "
        "SPA garantit une expérience utilisateur fluide, sans rafraîchissement complet des pages à chaque action."
        "\n- Le Tier Logique (Backend) : Un serveur Spring Boot exposé sous forme d'API RESTful. Il traite la logique métier, "
        "gère l'authentification, valide les entrées et filtre les accès sécurisés aux ressources privées."
        "\n- Le Tier de Persistance (Database) : Une base de données relationnelle PostgreSQL hébergée de manière managée "
        "sur la plateforme de cloud Neon DB, assurant des transactions ACID et une scalabilité à la demande."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "2. Justification du Choix du Backend : Spring Boot (Java 17)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour le développement du backend, le framework Spring Boot (Java 17) a été retenu pour ses performances "
        "en environnement d'entreprise, sa robustesse et ses capacités d'intégration de sécurité de haut niveau. "
        "Les modules spécifiques utilisés incluent :"
        "\n- Spring Boot Web : Pour exposer l'API RESTful rapidement avec un serveur Tomcat embarqué."
        "\n- Spring Data JPA : Simplifie l'interaction avec le SGBD PostgreSQL grâce au mapping objet-relationnel (ORM Hibernate). "
        "Il permet de réduire le code chaudière lié aux transactions et requêtes SQL manuelles."
        "\n- Spring Security : Fournit un framework d'authentification et de contrôle d'accès hautement personnalisable. "
        "Dans notre architecture, nous avons surchargé la configuration par défaut pour mettre en place une authentification "
        "stateless basée sur des jetons JWT, protégeant ainsi efficacement les routes d'écriture de l'API (création, édition et suppression de produits)."
        "\nL'utilisation de Java 17 apporte également de meilleures performances d'exécution, une gestion optimisée de la mémoire et la syntaxe moderne des Records pour les DTOs."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Justification du Choix du Frontend : React (Vite, Tailwind)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le client web a été conçu avec la bibliothèque React et l'outil de build ultra-rapide Vite. "
        "Les arguments en faveur de cette stack sont :"
        "\n- Performance au Build et au Dev : Vite utilise les ES Modules natifs du navigateur pour offrir un rechargement à chaud (HMR) "
        "quasi-instantané lors du développement, évitant les lourdeurs de Webpack."
        "\n- Rendu Composant React : Permet de concevoir des composants modulaires et réutilisables (ProductCard, Navbar, FavoriteButton) "
        "favorisant la maintenabilité du code."
        "\n- Gestion Globale de l'État : Mise en œuvre du React Context API (via MarketplaceContext) pour stocker les informations de session "
        "utilisateur (jeton JWT, informations de profil) et la liste partagée des produits, simplifiant les flux de données inter-composants."
        "\n- Rapidité de chargement : L'intégration de Tailwind CSS permet de styliser les pages rapidement tout en générant un bundle CSS final minime. "
        "Tailwind CSS permet de concevoir une interface responsive de manière déclarative à l'aide de classes utilitaires performantes."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "4. Justification de la Base de Données : PostgreSQL Serverless (Neon DB)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "PostgreSQL a été sélectionné pour sa conformité totale aux normes ACID, garantissant l'intégrité absolue "
        "des transactions (comme la relation d'achat, les associations de favoris, ou l'unicité des comptes d'utilisateurs). "
        "Le choix de Neon comme fournisseur de base de données managée apporte plusieurs avantages concurrentiels :"
        "\n- Architecture Serverless : Neon sépare le stockage du calcul, permettant à la base de données de s'éteindre automatiquement "
        "en période d'inactivité et de redémarrer en moins de 2 secondes lors d'un appel API. Cela réduit considérablement les coûts d'infrastructure "
        "pendant les heures de faible activité."
        "\n- Base de données Cloud-Native : Connexions sécurisées intégrées et console d'administration interactive permettant d'analyser "
        "les requêtes lentes et les schémas directement depuis un navigateur web."
        "\n- Branchement instantané : Permet de cloner la base de données de production en un clic pour tester des migrations ou des nouveautés sans risque."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "5. Containerisation et Infrastructure DevOps de Développement", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour garantir la portabilité de l'application entre les environnements de développement, de test et de production, "
        "une stratégie de conteneurisation intégrale avec Docker a été mise en œuvre. Le backend est compilé dans un conteneur standard "
        "reposant sur une image JDK 17 alpine ultra-légère. L'utilisation de Docker Compose simplifie le lancement local de l'application "
        "en orchestrant les conteneurs nécessaires avec des configurations standardisées pour les variables d'environnement et l'exposition des ports."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 4
    add_heading_styled(doc, "CONCEPTION DU SYSTÈME & MODÉLISATION UML", level=1)
    
    add_heading_styled(doc, "1. Modélisation Conceptuelle de Données (Diagramme de Classes Métier)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le diagramme de classes métier (ou modèle de domaine) modélise les concepts clés du problème B2B, "
        "sans faire référence aux détails d'implémentation technologiques. Il met en évidence trois concepts clés :"
        "\n- User : Représente tout professionnel enregistré sur la plateforme. Il possède un identifiant unique, un nom, "
        "un e-mail (qui fait office d'identifiant unique de connexion) et un numéro de téléphone WhatsApp pour les prises de contact."
        "\n- Produit : Représente une annonce de produit/matériel publiée sur la plateforme. Chaque produit possède un prix, "
        "un statut de vente (Disponible ou Vendu), une image, une date de publication, une marque et un compteur de vues."
        "\n- CategorieProduit : Sert à classifier logiquement les produits pour faciliter la navigation interentreprises."
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "business_class_diagram.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 4.1 : Diagramme de Classes Métier (Domain Model)")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)

    add_heading_styled(doc, "2. Spécification Technique des Entités (Diagramme de Classes Technique)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le diagramme de classes technique détaille les relations structurelles au niveau de l'implémentation Java. "
        "Il intègre les différentes couches de l'architecture logicielle : les Contrôleurs REST qui interceptent les requêtes HTTP, "
        "les Services qui orchestrent la logique métier, les Repositories (Spring Data) responsables des requêtes de persistance, "
        "et les Entités JPA mappées à notre base de données PostgreSQL Neon :"
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "technical_class_diagram.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 4.2 : Diagramme de Classes Technique (Architecture en Couches Backend)")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)

    add_heading_styled(doc, "3. Dictionnaire de Données et Structure des Tables SQL", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Les relations UML sont transposées au niveau de la base de données relationnelle sous forme de tables SQL structurées "
        "de la façon suivante :"
        "\n- Table `users` : Contient les informations des utilisateurs. Un index unique est appliqué sur l'e-mail pour garantir son unicité physique."
        "\n- Table `categories_produit` : Contient le catalogue de catégorisation."
        "\n- Table `produits` : Contient le catalogue d'annonces. Une clé étrangère `vendeur_id` pointe vers la table `users` et une clé `categorie_id` pointe vers la table `categories_produit`."
        "\n- Table d'association `user_favoris` : Table de jointure modélisant la relation many-to-many entre `users` (colonne `user_id`) et `produits` (colonne `produit_id`)."
        "\nL'utilisation d'index sur les clés étrangères (`vendeur_id`, `categorie_id`) permet d'optimiser les performances des requêtes de jointure "
        "lorsque le volume d'annonces augmente."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 5
    add_heading_styled(doc, "MODÈLE DE SÉCURITÉ ET AUTHENTIFICATION JWT", level=1)
    
    add_heading_styled(doc, "1. Principes Fondamentaux de la Sécurité Stateless avec JWT", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour assurer une architecture RESTful évolutive et indépendante (stateless), la plateforme utilise des jetons JWT "
        "(JSON Web Tokens) pour sécuriser ses échanges. L'utilisation du JWT évite l'utilisation de sessions HTTP côté serveur "
        "qui consommeraient trop de ressources en mémoire vive. Un jeton JWT est composé de trois parties distinctes séparées "
        "par des points : le Header (indiquant l'algorithme de hachage), le Payload (contenant les revendications ou claims comme l'ID utilisateur "
        "et la date d'expiration) et la Signature (calculée à l'aide d'une clé secrète serveur pour interdire toute falsification). "
        "Le client stocke le jeton et le transmet dans l'en-tête `Authorization` à chaque requête de modification de données."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "2. Implémentation du Filtre de Sécurité Customisé Spring Boot", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Dans le backend Spring Boot, la sécurité est mise en place par le composant `SecurityConfig` et un filtre de servlet personnalisé "
        "nommé `JwtAuthFilter` (qui hérite de `OncePerRequestFilter`). Le flux de sécurité se déroule ainsi :"
        "\n1. L'application cliente (React) envoie une requête HTTP vers une route protégée, en incluant l'en-tête standard : `Authorization: Bearer <token_jwt>`."
        "\n2. Le filtre `JwtAuthFilter` intercepte la requête HTTP avant qu'elle n'atteigne le contrôleur Spring Web."
        "\n3. Le filtre extrait le préfixe 'Bearer ' et analyse le jeton à l'aide du composant utilitaire `JwtUtil`."
        "\n4. Si le jeton est valide, l'ID utilisateur est extrait et injecté sous forme d'objet d'authentification principal dans le contexte de sécurité de Spring Security (`SecurityContextHolder`)."
        "\n5. Le contrôleur REST correspondant peut alors récupérer directement l'identifiant de l'utilisateur connecté via l'annotation `@AuthenticationPrincipal Long userId`."
        "\n6. Si le jeton est manquant, corrompu ou expiré, la requête est rejetée avec un code HTTP 401 Unauthorized."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Gestion de la Politique de Partage des Ressources (CORS)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Puisque le frontend React et le backend Spring Boot sont hébergés sur des serveurs et des domaines distincts (Vercel et AWS EC2), "
        "le navigateur applique la politique de sécurité de même origine (Same-Origin Policy). Pour permettre les requêtes cross-origin, "
        "le composant `CorsConfig` du backend Spring Boot configure explicitement les règles autorisées :"
        "\n- Origines autorisées : Les URL spécifiques du frontend hébergé sur Vercel (ainsi que `http://localhost:5173` pour le développement local)."
        "\n- Méthodes HTTP autorisées : GET, POST, PUT, PATCH, DELETE et OPTIONS."
        "\n- En-têtes autorisés : Content-Type, Authorization, et tous les en-têtes personnalisés."
        "\n- Credentials : Les cookies ou en-têtes d'authentification sont autorisés à transiter en toute sécurité."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 6
    add_heading_styled(doc, "CONCEPTION DÉTAILLÉE DES FLUX ET SÉQUENCES", level=1)
    
    add_heading_styled(doc, "1. Diagramme de Séquence : Publication d'un Produit par le Vendeur", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Ce diagramme de séquence montre l'ensemble des échanges réseau et des opérations logiques internes requis "
        "lorsqu'un vendeur connecté décide de publier une nouvelle annonce de produit sur la plateforme. "
        "Ce flux prend en charge le traitement d'une requête multipart pour le téléversement de la photo du produit :"
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "sequence_add_product.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 6.1 : Diagramme de Séquence - Publication d'un Produit")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)

    add_heading_styled(doc, "2. Diagramme de Séquence : Prise de Contact Directe via WhatsApp", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La mise en relation directe est une fonctionnalité essentielle de notre marketplace B2B. Le diagramme suivant "
        "représente le processus d'interaction où un acheteur (authentifié ou non) consulte la fiche d'un produit, "
        "ce qui incrémente le compteur de vues de l'annonce côté backend, puis clique sur le bouton de contact WhatsApp "
        "pour entamer instantanément une conversation de négociation :"
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "sequence_contact_seller.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 6.2 : Diagramme de Séquence - Contact Vendeur via WhatsApp")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)

    add_heading_styled(doc, "3. Diagramme de Séquence : Partage de l'Annonce avec un Ami", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour favoriser le partage viral des opportunités d'affaires (B2B), nous avons spécifié un mécanisme de partage "
        "intégré. Ce diagramme décrit le cas d'utilisation où un professionnel partage une annonce active avec un collègue "
        "ou un ami via les fonctionnalités de son appareil en s'appuyant sur l'API Web Share native :"
    )
    format_para(p, space_after=10)
    
    # Insert diagram image
    doc.add_picture(os.path.join(IMAGES_DIR, "sequence_share_ad.png"), width=Inches(6.0))
    p_cap = doc.add_paragraph()
    p_cap_run = p_cap.add_run("Figure 6.3 : Diagramme de Séquence - Partage d'une Publication")
    format_para(p_cap, size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=18)

    add_heading_styled(doc, "4. Implémentation du Partage Frontend (Web Share API et Presse-Papiers)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour assurer le partage d'annonces de manière professionnelle et fluide, le code frontend React "
        "déploie une fonction de partage dynamique. Si le navigateur prend en charge l'API Web Share de type "
        "`navigator.share()`, l'application propose l'ouverture du volet de partage système de l'OS (permettant un partage "
        "vers des applications comme WhatsApp, Slack, Teams, Gmail, ou par SMS). "
        "En guise de repli (fallback) pour les navigateurs non compatibles, le système copie instantanément le lien court de "
        "l'annonce dans le presse-papiers de l'utilisateur tout en affichant un toast de notification élégant indiquant "
        "'Lien copié dans le presse-papiers !'. Voici le principe de cette implémentation :"
    )
    format_para(p, space_after=12)
    
    add_callout(doc,
                "const handleShare = async () => {\n"
                "  const shareData = {\n"
                "    title: product.name,\n"
                "    text: `Consultez ce produit sur notre marketplace B2B : ${product.name}`,\n"
                "    url: window.location.href\n"
                "  };\n"
                "  if (navigator.share) {\n"
                "    try {\n"
                "      await navigator.share(shareData);\n"
                "    } catch (err) {\n"
                "      console.log('Erreur lors du partage natif:', err);\n"
                "    }\n"
                "  } else {\n"
                "    navigator.clipboard.writeText(window.location.href);\n"
                "    showToast('Lien de l\\'annonce copié !');\n"
                "  }\n"
                "};",
                "CODE JS: LOGIQUE DE PARTAGE D'ANNONCE", "1B365D")

    doc.add_page_break()

    # CHAPITRE 7
    add_heading_styled(doc, "DEVOPS, DOCKER & WORKFLOW DE DÉPLOIEMENT", level=1)
    
    add_heading_styled(doc, "1. Configuration des Conteneurs via Dockerfile Multi-stage", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Dans une démarche moderne d'automatisation du déploiement, nous utilisons un workflow basé sur Docker. "
        "Le backend Spring Boot possède un fichier `Dockerfile` multi-stage permettant de compiler puis de packager "
        "l'application de manière optimale, en limitant la taille de l'image de production finale. La compilation est effectuée "
        "dans un premier environnement Maven temporaire, et seul le jar final est copié dans l'environnement d'exécution."
    )
    format_para(p, space_after=12)
    
    # Styled code block for Dockerfile
    add_callout(doc, 
                "# Stage 1: Build the Maven application\n"
                "FROM maven:3.8.5-openjdk-17 AS builder\n"
                "WORKDIR /app\n"
                "COPY pom.xml .\n"
                "COPY src ./src\n"
                "RUN mvn clean package -DskipTests\n\n"
                "# Stage 2: Create runtime image\n"
                "FROM openjdk:17-jdk-alpine\n"
                "WORKDIR /app\n"
                "COPY --from=builder /app/target/b2b-0.0.1-SNAPSHOT.jar app.jar\n"
                "EXPOSE 8080\n"
                "ENTRYPOINT [\"java\", \"-jar\", \"app.jar\"]",
                "DOCKERFILE BACKEND (Spring Boot)", "2C3E50")

    add_heading_styled(doc, "2. Orchestration Locale Multiconteneur avec Docker Compose", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour exécuter l'infrastructure complète localement lors des phases de développement ou de test, "
        "un fichier `compose.yaml` est mis en place dans le répertoire backend. Ce fichier orchestre une base de données "
        "PostgreSQL locale ainsi que l'application Spring Boot conteneurisée. En production, la base PostgreSQL locale est "
        "substituée par la base PostgreSQL serverless Neon DB pour de meilleures garanties de résilience, de sauvegardes automatiques et de performance. "
        "Cela garantit que l'équipe de développement travaille dans des conditions proches de la production."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 8
    add_heading_styled(doc, "ARCHITECTURE DE DÉPLOIEMENT CLOUD", level=1)
    
    add_heading_styled(doc, "1. Vue d'Ensemble de l'Architecture Déployée", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'application est déployée sur trois services cloud distincts, chacun hébergeant une couche spécifique de l'architecture :"
        "\n- Le backend Spring Boot tourne sur une instance AWS EC2 (machine virtuelle Linux dans le cloud d'Amazon)."
        "\n- Le frontend React est hébergé sur Vercel, plateforme spécialisée dans l'hébergement d'applications web statiques."
        "\n- La base de données PostgreSQL est hébergée sur Neon, un service de base de données cloud managé."
        "\nCes trois composants communiquent entre eux via le protocole HTTPS. Le frontend envoie des requêtes HTTP "
        "vers l'API REST exposée par le backend EC2. Le backend, quant à lui, se connecte à la base de données Neon "
        "via une URL de connexion sécurisée (SSL/TLS)."
    )
    format_para(p, space_after=12)

    add_callout(doc,
                "Frontend (Vercel) ──HTTPS──▶ API REST (AWS EC2 / Spring Boot) ──JDBC/SSL──▶ Base de données (Neon / PostgreSQL)",
                "FLUX DE COMMUNICATION ENTRE COMPOSANTS", "2C3E50")

    add_heading_styled(doc, "2. Hébergement du Backend sur AWS EC2", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le backend Spring Boot est déployé sous forme de conteneur Docker sur une instance Linux AWS EC2. "
        "Un serveur web Nginx fait office de point d'entrée : il reçoit les requêtes HTTPS arrivant sur le domaine public "
        "et les redirige vers l'application Spring Boot en interne. "
        "Les groupes de sécurité AWS restreignent l'accès réseau : seul le trafic HTTP/HTTPS public est autorisé, "
        "le port interne de Spring Boot n'est pas exposé directement sur internet."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Hébergement du Frontend sur Vercel", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le client React est compilé (`npm run build`) en un ensemble de fichiers statiques (HTML, JS, CSS) "
        "puis déployé sur Vercel via une intégration GitHub. À chaque mise à jour du code sur le dépôt, "
        "Vercel relance automatiquement un build et publie la nouvelle version."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "4. Communication entre les Composants", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le frontend communique avec le backend via des requêtes HTTP REST. L'URL de base de l'API "
        "(`VITE_API_URL`) est configurée comme variable d'environnement dans Vercel et pointe vers le domaine "
        "public de l'instance EC2. Toutes les communications transitent en HTTPS pour garantir la confidentialité "
        "des données (en particulier les jetons JWT et les informations utilisateurs). "
        "\nLa base de données Neon est accessible par le backend uniquement, via une URL de connexion JDBC sécurisée. "
        "Le frontend n'a jamais de connexion directe à la base de données : toutes les opérations passent obligatoirement "
        "par l'API REST du backend, qui valide les droits d'accès avant chaque action."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 9
    add_heading_styled(doc, "GUIDE D'UTILISATION COMPLET DE LA PLATEFORME B2B", level=1)
    
    add_heading_styled(doc, "1. Inscription, Connexion et Gestion du Profil Utilisateur", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Pour tirer pleinement parti des fonctionnalités de la plateforme B2B, l'utilisateur doit disposer d'un compte actif. "
        "La procédure s'articule autour des étapes suivantes :"
        "\n- Inscription : L'utilisateur se rend sur la page d'inscription. Il saisit son nom de société ou d'individu, "
        "son adresse e-mail professionnelle unique, son mot de passe (qui doit comporter au moins 6 caractères pour des raisons de sécurité) "
        "et son numéro de téléphone WhatsApp au format international (ex: +22890000000). Le clic sur le bouton d'inscription soumet les données "
        "et connecte automatiquement l'utilisateur en générant un jeton de session JWT."
        "\n- Connexion : Si l'utilisateur possède déjà un compte, il renseigne son adresse e-mail et son mot de passe sur la page de connexion."
        "\n- Suivi de Session : L'application conserve les informations de l'utilisateur connecté en haut à droite de l'écran dans un menu de profil personnalisé. "
        "Un bouton de déconnexion permet de détruire le jeton stocké dans le navigateur pour sécuriser son poste de travail."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "2. Recherche Avancée, Filtres par Catégorie et Liste de Favoris", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La page d'accueil (Accueil.jsx) est conçue pour être claire et intuitive. Elle présente un en-tête (Hero section) invitant "
        "les professionnels à explorer les offres industrielles et matérielles actives. "
        "\n- Liste des Produits : Les annonces s'affichent sous forme de grille de cartes produit. Chaque carte indique si le produit est 'Nouveau', "
        "son prix, sa catégorie et un aperçu de son image."
        "\n- Filtres Rapides : Des boutons de filtre par catégorie permettent en un clic d'isoler les matériels informatiques ou les fournitures de bureau."
        "\n- Barre de Recherche : L'utilisateur peut saisir des mots-clés dans le champ de recherche pour filtrer dynamiquement les annonces."
        "\n- Liste des Favoris : Les acheteurs connectés peuvent cliquer sur le cœur d'une carte produit pour l'ajouter aux favoris. La liste est accessible "
        "dans l'onglet 'Favoris' (Favorites.jsx)."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "3. Espace Vendeur : Ajout, Édition et Clôture d'une Annonce", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Tout utilisateur authentifié a accès au rôle de Vendeur et dispose de son propre Tableau de Bord (Dashboard.jsx) :"
        "\n- Publier une annonce : Le vendeur clique sur 'Ajouter un produit'. Un formulaire complet lui permet de renseigner le nom, "
        "la description technique, le prix, la marque, le numéro WhatsApp et de téléverser l'image du produit. À la soumission, le produit est publié immédiatement."
        "\n- Modifier une annonce : Le vendeur peut corriger des erreurs de prix, de description ou téléverser une nouvelle photo à tout moment."
        "\n- Clôturer/Marquer comme Vendu : Le vendeur peut basculer le statut du produit de 'Disponible' à 'Vendu'. Un badge rouge 'VENDU' s'affichera alors sur l'annonce, "
        "et le bouton WhatsApp sera automatiquement désactivé pour éviter de nouvelles sollicitations."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "4. Utilisation Pratique des Boutons de Contact et de Partage d'Annonces", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La fonctionnalité de partage d'annonce a été spécifiquement conçue pour encourager la diffusion collaborative des offres :"
        "\n1. Accédez à la page de détails de l'annonce que vous souhaitez partager."
        "\n2. Repérez le bouton 'Partager' situé à côté des informations de contact."
        "\n3. Clic sur le bouton : si vous naviguez depuis un smartphone, un menu système s'ouvre pour envoyer l'annonce directement "
        "à un ami sur WhatsApp, Messenger, par e-mail ou Slack."
        "\n4. Si le partage natif n'est pas disponible (sur PC), l'application copie automatiquement l'URL directe de l'annonce "
        "dans le presse-papiers, et vous pouvez la coller où vous le souhaitez. De plus, le bouton WhatsApp permet d'entamer une conversation "
        "directe avec le vendeur avec un texte prérempli."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 10
    add_heading_styled(doc, "ADMINISTRATION ET MAINTENANCE DE LA PLATEFORME", level=1)
    
    add_heading_styled(doc, "1. Gestion de la Base de Données via la Console Neon", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La base de données PostgreSQL hébergée sur Neon est administrable directement via sa console web. "
        "Cette interface permet d'exécuter des requêtes SQL pour vérifier les données stockées "
        "(liste des utilisateurs, état des annonces, associations de favoris), "
        "de consulter les tables et leur schéma, et de modifier des enregistrements si nécessaire. "
        "Neon propose également un système de branches de base de données, permettant de tester des "
        "modifications de schéma sur une copie isolée sans affecter les données de production."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "2. Supervision de l'Application Backend", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Le backend Spring Boot s'exécutant dans un conteneur Docker sur l'instance EC2, "
        "les journaux applicatifs (logs) sont accessibles via la commande Docker standard (`docker logs`). "
        "Ces logs permettent de diagnostiquer les erreurs éventuelles : échecs d'authentification JWT, "
        "problèmes de connexion à la base de données, ou requêtes REST malformées. "
        "\nLes logs Nginx enregistrent toutes les requêtes HTTP entrantes (méthode, URL, code de réponse), "
        "ce qui permet de vérifier que le routage entre le frontend Vercel et le backend EC2 fonctionne correctement."
    )
    format_para(p, space_after=12)

    doc.add_page_break()

    # CHAPITRE 11
    add_heading_styled(doc, "VALIDATION ET TESTS DE L'APPLICATION", level=1)
    
    add_heading_styled(doc, "1. Tests Fonctionnels et Vérification des Scénarios Clés", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "La phase de validation de notre projet a consisté à vérifier manuellement et de façon systématique le bon "
        "déroulement de chacun des scénarios définis dans le cahier des charges. Pour chaque cas d'utilisation principal, "
        "nous avons élaboré un protocole de test qui spécifie les données d'entrée, les actions à effectuer et le résultat attendu. "
        "\nVoici les principaux scénarios validés :"
        "\n- Inscription / Connexion : Création d'un compte avec des données valides puis avec un e-mail déjà utilisé pour vérifier la gestion d'erreur."
        "\n- Publication d'annonce : Soumission du formulaire complet avec téléversement d'image. Vérification de l'apparition immédiate du produit dans le catalogue."
        "\n- Sécurité JWT : Tentative d'appel direct à l'API REST de création de produit sans jeton d'authentification — "
        "vérification du retour HTTP 401 Unauthorized."
        "\n- Contact WhatsApp : Clic sur le bouton 'Contacter le Vendeur' depuis la fiche produit — vérification que l'application WhatsApp s'ouvre "
        "avec le bon numéro et le texte prérempli."
        "\n- Partage d'annonce : Test du bouton de partage sur mobile (Android) — vérification de l'affichage du volet natif du système. "
        "Test sur navigateur de bureau — vérification que l'URL est bien copiée dans le presse-papiers."
        "\n- Gestion des favoris : Ajout et retrait d'un produit depuis la liste de favoris, avec vérification de la mise à jour visuelle immédiate."
    )
    format_para(p, space_after=12)

    add_heading_styled(doc, "2. Résultats et Points d'Amélioration", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "L'ensemble des fonctionnalités du cahier des charges a été implémenté et validé. "
        "Quelques points méritent cependant d'être améliorés dans une version ultérieure :"
        "\n- La gestion CORS entre Vercel et EC2 nécessite une configuration explicite des origines autorisées "
        "dans la classe `CorsConfig` de Spring Boot pour que le frontend puisse appeler l'API sans erreurs."
        "\n- Le téléversement d'images (stockées en `byte[]` dans PostgreSQL) requiert une gestion soigneuse "
        "du format `multipart/form-data` côté backend et de l'objet `FormData` côté React."
        "\n- La compatibilité de l'API Web Share : `navigator.share()` n'est pas disponible sur tous les navigateurs. "
        "Un mécanisme de repli (`fallback`) copie l'URL dans le presse-papiers lorsque le partage natif est indisponible."
    )
    format_para(p, space_after=12)

    add_callout(doc,
                "L'ensemble des fonctionnalités du cahier des charges initial a été implémenté et validé avec succès. "
                "Les tests ont été conduits sur Google Chrome, Firefox et sur mobile Android. "
                "Aucun défaut bloquant n'a été identifié lors de la validation finale.",
                "RÉSULTAT DE LA VALIDATION", "16A085")

    add_heading_styled(doc, "CONCLUSION ET PERSPECTIVES", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "La plateforme marketplace B2B réalisée dans le cadre de ce cours de Conception des Systèmes d'Information "
        "constitue une application web fullstack fonctionnelle et déployée. "
        "Elle couvre l'ensemble des besoins définis dans le cahier des charges : publication d'annonces, "
        "recherche et filtrage, gestion des favoris, contact direct via WhatsApp et partage d'annonces. "
        "L'architecture adoptée — Spring Boot, React, PostgreSQL Neon, Docker et déploiement cloud — "
        "reflète les bonnes pratiques actuelles du développement logiciel."
    )
    format_para(p, space_after=10)

    add_heading_styled(doc, "Bilan Technique du Projet", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Sur le plan technique, le projet a permis de mettre en œuvre les concepts suivants :"
        "\n- Architecture multicouche (Controller, Service, Repository, Entity) avec Spring Boot."
        "\n- Développement d'une Single Page Application en React avec gestion de l'état global via Context API."
        "\n- Authentification stateless par JSON Web Tokens (JWT) avec filtre de sécurité Spring Security."
        "\n- Modélisation UML complète : diagramme de cas d'utilisation, diagramme de classes métier "
        "et technique, diagrammes de séquences pour les flux principaux."
        "\n- Containerisation avec Docker et déploiement distribué sur AWS EC2, Vercel et Neon DB."
        "\n- Intégration de la Web Share API et du contact WhatsApp pour une expérience utilisateur moderne."
    )
    format_para(p, space_after=10)

    add_heading_styled(doc, "Perspectives d'Amélioration", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Des évolutions peuvent être envisagées pour enrichir la plateforme :"
        "\n- Un système de messagerie interne en temps réel (WebSocket) pour permettre la négociation "
        "directement sur la plateforme."
        "\n- Un moteur de recommandation basé sur l'historique de navigation et les favoris."
        "\n- Un backoffice d'administration permettant de modérer les annonces et gérer les utilisateurs."
        "\n- Des notifications push pour alerter un vendeur lors d'une prise de contact."
        "\nLa plateforme développée présente un potentiel réel et pourrait, avec ces améliorations progressives, "
        "constituer une solution utilisable dans un contexte professionnel."
    )
    format_para(p, space_after=12)

    # SAUVEGARDE
    doc.save(OUTPUT_DOCX)
    print(f"Rapport de projet généré avec succès dans: {OUTPUT_DOCX}")

if __name__ == "__main__":
    print("Démarrage de la génération des diagrammes et du rapport...")
    generate_diagrams()
    generate_report()
    print("Terminé avec succès !")
